from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Deployment-and-Setup-Guide.docx"


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def format_run(run, *, name="Calibri", size=11, bold=False, color="000000"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)


def style_paragraph(paragraph, *, after=6, before=0, line=1.25, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    paragraph.alignment = alignment


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    style_paragraph(p, after=4, line=1.25)
    run = p.add_run(text)
    format_run(run)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    style_paragraph(p, after=4, line=1.25)
    run = p.add_run(text)
    format_run(run)
    return p


def add_heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    style_paragraph(
        p,
        before={1: 18, 2: 14, 3: 10}[level],
        after={1: 10, 2: 7, 3: 5}[level],
        line=1.15,
    )
    run = p.add_run(text)
    format_run(
        run,
        size={1: 16, 2: 13, 3: 12}[level],
        bold=True,
        color={1: "2E74B5", 2: "2E74B5", 3: "1F4D78"}[level],
    )
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    style_paragraph(p, after=6, line=1.25)
    run = p.add_run(text)
    format_run(run)
    return p


def add_package_table(doc, title, rows):
    add_heading(doc, title, 2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    widths = [Inches(2.1), Inches(1.8), Inches(2.6)]
    for idx, width in enumerate(widths):
        table.columns[idx].width = width
    headers = ["Package", "Version", "Purpose"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = widths[idx]
        cell.text = ""
        para = cell.paragraphs[0]
        style_paragraph(para, after=0, line=1.0)
        run = para.add_run(text)
        format_run(run, bold=True)
        set_cell_shading(cell, "E8EEF5")
        set_cell_margins(cell)
        set_cell_border(
            cell,
            top={"val": "single", "sz": 8, "color": "B7C5D6"},
            left={"val": "single", "sz": 8, "color": "B7C5D6"},
            bottom={"val": "single", "sz": 8, "color": "B7C5D6"},
            right={"val": "single", "sz": 8, "color": "B7C5D6"},
        )
    for package, version, purpose in rows:
        row = table.add_row()
        for idx, value in enumerate((package, version, purpose)):
            cell = row.cells[idx]
            cell.width = widths[idx]
            cell.text = ""
            para = cell.paragraphs[0]
            style_paragraph(para, after=0, line=1.15)
            run = para.add_run(value)
            format_run(run)
            set_cell_margins(cell)
            set_cell_border(
                cell,
                top={"val": "single", "sz": 6, "color": "DADCE0"},
                left={"val": "single", "sz": 6, "color": "DADCE0"},
                bottom={"val": "single", "sz": 6, "color": "DADCE0"},
                right={"val": "single", "sz": 6, "color": "DADCE0"},
            )
    doc.add_paragraph()


def add_footer(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Deployment and Setup Guide")
    format_run(run, size=9, color="666666")


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.line_spacing = 1.0
    run = title.add_run("Deployment and Setup Guide")
    format_run(run, size=22, bold=True, color="2E74B5")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.add_run("Project checklist for installation, packages, configuration, and GitHub upload")
    format_run(subtitle_run, size=11, color="555555")

    add_body(
        doc,
        "This document explains what software is needed, what packages are used by the project, "
        "how to run the backend and frontend on a fresh machine, and what to verify before uploading the project to GitHub.",
    )

    add_heading(doc, "1. Required Software", 1)
    for item in [
        "Git",
        "Python 3.11 or newer",
        "Node.js 20 or newer",
        "npm",
        "Docker Desktop if PostgreSQL mode will be used",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2. Backend Python Packages", 1)
    add_body(doc, "These packages are listed in backend/requirements.txt and are required for the FastAPI backend.")
    add_package_table(
        doc,
        "Backend package list",
        [
            ("fastapi", ">=0.116,<1.0", "Backend API framework"),
            ("uvicorn[standard]", ">=0.35,<1.0", "ASGI server for running the backend"),
            ("sqlalchemy", ">=2.0,<3.0", "Database ORM and query layer"),
            ("psycopg[binary]", ">=3.2.13,<4.0", "PostgreSQL database driver"),
            ("pydantic", ">=2.12,<3.0", "Schema and data validation"),
            ("pydantic-settings", ">=2.10,<3.0", "Environment-based configuration"),
            ("python-multipart", ">=0.0.20,<1.0", "Upload handling for multipart form data"),
        ],
    )

    add_heading(doc, "3. Frontend Node Packages", 1)
    add_body(doc, "These packages are listed in frontend/package.json and are needed for the Next.js frontend.")
    add_package_table(
        doc,
        "Frontend package list",
        [
            ("next", "15.3.5", "Frontend framework"),
            ("react", "19.1.0", "UI library"),
            ("react-dom", "19.1.0", "Browser rendering for React"),
            ("typescript", "5.8.3", "Type checking and development tooling"),
            ("@types/node", "24.0.10", "TypeScript types for Node.js"),
            ("@types/react", "19.1.8", "TypeScript types for React"),
            ("@types/react-dom", "19.1.6", "TypeScript types for React DOM"),
        ],
    )

    add_heading(doc, "4. Database and Deployment Requirements", 1)
    add_bullet(doc, "The project supports PostgreSQL and also supports a local SQLite fallback.")
    add_bullet(doc, "docker-compose.yml uses the postgres:16 Docker image.")
    add_bullet(doc, "Database-related Python packages are sqlalchemy and psycopg[binary].")

    add_heading(doc, "5. Environment Values To Verify", 1)
    add_heading(doc, "Backend values", 2)
    for item in [
        "DATABASE_URL",
        "POSTGRES_HOST",
        "POSTGRES_PORT",
        "POSTGRES_DB",
        "POSTGRES_USER",
        "POSTGRES_PASSWORD",
        "USE_SQLITE_FALLBACK",
        "SQLITE_FALLBACK_URL",
        "REFERENCE_ACCESSION",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Frontend values", 2)
    for item in [
        "NEXT_PUBLIC_API_BASE_URL",
        "BACKEND_API_BASE_URL",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Recommended local defaults", 2)
    add_bullet(doc, "NEXT_PUBLIC_API_BASE_URL=/api/backend")
    add_bullet(doc, "BACKEND_API_BASE_URL=http://127.0.0.1:8000/api/v1")

    add_heading(doc, "6. Backend Setup Steps", 1)
    for step in [
        "Open a terminal in the backend folder.",
        "Create a virtual environment: python -m venv .venv",
        "Activate it: .venv\\Scripts\\Activate.ps1",
        "Upgrade pip: python -m pip install --upgrade pip",
        "Install backend packages: pip install -r requirements.txt",
        "Initialize the database: python -m app.init_db",
        "Run the backend server: python -m uvicorn app.main:app --reload --host 127.0.0.1 --port 8000",
    ]:
        add_number(doc, step)
    add_heading(doc, "Backend checks", 2)
    for item in [
        "http://127.0.0.1:8000/health",
        "http://127.0.0.1:8000/health/database",
        "http://127.0.0.1:8000/docs",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "7. Frontend Setup Steps", 1)
    for step in [
        "Open a terminal in the frontend folder.",
        "Install frontend packages: npm install",
        "Run the frontend server: npm run dev",
        "Open the app in the browser at http://127.0.0.1:3000",
    ]:
        add_number(doc, step)

    add_heading(doc, "8. PostgreSQL Setup With Docker", 1)
    add_number(doc, "From the project root, run: docker compose up -d postgres")
    add_number(doc, "Use DATABASE_URL=postgresql+psycopg://postgres:postgres@localhost:5432/sarscov2")
    add_number(doc, "Run python -m app.init_db after PostgreSQL starts")

    add_heading(doc, "9. SQLite Fallback Mode", 1)
    add_bullet(doc, "Use USE_SQLITE_FALLBACK=true")
    add_bullet(doc, "Use SQLITE_FALLBACK_URL=sqlite:///./sarscov2.db")
    add_bullet(doc, "This is the easiest mode for first local testing when PostgreSQL is not ready.")

    add_heading(doc, "10. How To Check Packages Yourself", 1)
    add_bullet(doc, "Check backend packages in backend/requirements.txt")
    add_bullet(doc, "Check frontend packages in frontend/package.json")
    add_bullet(doc, "Use pip list inside the backend virtual environment to see installed Python packages")
    add_bullet(doc, "Use npm list --depth=0 inside frontend to see installed Node packages")
    add_bullet(doc, "Use code imports to confirm what the project actually uses")

    add_heading(doc, "11. GitHub Upload Checklist", 1)
    add_heading(doc, "Files to include", 2)
    for item in [
        "backend/",
        "frontend/",
        "docs/",
        "docker-compose.yml",
        "README.md",
        "backend/requirements.txt",
        "frontend/package.json",
        "frontend/package-lock.json",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Files not to upload", 2)
    for item in [
        "node_modules/",
        ".venv/",
        "__pycache__/",
        ".next/",
        ".env files",
        "*.db, *.db-shm, *.db-wal",
        "*.log",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "12. Final Verification Before Sharing", 1)
    for item in [
        "Backend starts without errors",
        "Frontend starts without errors",
        "Backend health endpoint returns OK",
        "Database health endpoint returns OK",
        "Frontend can reach the backend",
        "Upload flow works",
        "Jobs page loads",
        "Sample page loads",
    ]:
        add_bullet(doc, item)

    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
